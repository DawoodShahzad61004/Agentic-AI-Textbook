from __future__ import annotations

import re
import sys
from pathlib import Path
from zipfile import ZipFile

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / ".docx-tools"))

from docx import Document


FILES = {
    11: ROOT / "Chapters Created So Far" / "Chapter_11_Query_Retrieval_Fundamentals.docx",
    12: ROOT / "Chapters Created So Far" / "Chapter_12_Advanced_Retrieval_Techniques.docx",
    13: ROOT / "Chapters Created So Far" / "Chapter_13_Generating_Answers_with_an_LLM.docx",
    "13B": ROOT / "Chapters Created So Far" / "Chapter_13B_Centralized_LLM_Invocation_and_Error_Handling.docx",
    18: ROOT / "Chapters Created So Far" / "Chapter_18_Implementing_the_Agent.docx",
    19: ROOT / "Chapters Created So Far" / "Chapter_19_Orchestration_with_LangGraph.docx",
    "19B": ROOT / "Chapters Created So Far" / "Chapter_19B_Porting_the_Agent_to_LangGraph.docx",
    20: ROOT / "Chapters Created So Far" / "Chapter_20_Quality_Control_and_Self_Correction.docx",
    "20B": ROOT / "Chapters Created So Far" / "Chapter_20B_Structured_Output_Reliability.docx",
    21: ROOT / "Chapters Created So Far" / "Chapter_21_Answer_Relevance_Verification.docx",
    22: ROOT / "Chapters Created So Far" / "Chapter_22_Observability_Dry_Run_Trace.docx",
    "22B": ROOT / "Chapters Created So Far" / "Chapter_22B_Semantic_Compression.docx",
    "22C": ROOT / "Chapters Created So Far" / "Chapter_22C_Two_Track_Parallel_Compression.docx",
    14: ROOT / "Chapters Created So Far" / "Chapter_14_Prompt_Engineering_for_RAG.docx",
    "5B": ROOT / "Chapters Created So Far" / "Chapter_5B_Evaluating_Document_Conversion_Engines.docx",
    "7B": ROOT / "Chapters Created So Far" / "Chapter_7B_Chunking_Converted_Documents.docx",
    15: ROOT / "Chapters Created So Far" / "Chapter_15_What_Agentic_Really_Means.docx",
    16: ROOT / "Chapters Created So Far" / "Chapter_16_Designing_the_Agent_Loop.docx",
    17: ROOT / "Chapters Created So Far" / "Chapter_17_Tool_Use_and_Function_Calling.docx",
    28: ROOT / "Chapters Created So Far" / "Chapter_28_What_Self_Learning_Actually_Means.docx",
    40: ROOT / "Chapters Created So Far" / "Chapter_40_Security_and_Safety.docx",
}

EXPECTED = {
    11: ["11.1", "11.2", "11.3", "11.4", "11.5", "11.6", "11.7"],
    12: [f"12.{i}" for i in range(1, 11)],
    13: [f"13.{i}" for i in range(1, 10)],
    "13B": [f"13B.{i}" for i in range(1, 20)],
    18: [f"18.{i}" for i in range(1, 12)],
    19: [f"19.{i}" for i in range(1, 15)],
    "19B": [f"19B.{i}" for i in range(1, 15)],
    20: [f"20.{i}" for i in range(1, 12)],
    "20B": [f"20B.{i}" for i in range(1, 12)],
    21: [f"21.{i}" for i in range(1, 8)],
    22: [f"22.{i}" for i in range(1, 12)],
    "22B": [f"22B.{i}" for i in range(1, 12)],
    "22C": [f"22C.{i}" for i in range(1, 12)],
    14: [f"14.{i}" for i in range(1, 17)],
    "5B": [f"5B.{index}" for index in range(1, 11)],
    "7B": [f"7B.{index}" for index in range(1, 13)],
    15: ["15.1", "15.2", "15.3", "15.4", "15.5"],
    16: ["16.1", "16.2", "16.3", "16.4", "16.5"],
    17: ["17.1", "17.2", "17.3", "17.4", "17.5", "17.6", "17.7"],
    28: ["28.1", "28.2", "28.3", "28.4", "28.5"],
    40: ["40.1", "40.2", "40.3", "40.4", "40.5", "40.6"],
}

MIN_WORDS = {"5B": 2800, "7B": 3000}


def count_words(doc: Document) -> int:
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return len(re.findall(r"\b[\w’'-]+\b", "\n".join(chunks)))


def audit(chapter: int | str, path: Path) -> list[str]:
    errors: list[str] = []
    doc = Document(path)
    sec = doc.sections[0]
    expected = {
        "page_width": 8.5,
        "page_height": 11.0,
        "top_margin": 2.5 / 2.54,
        "left_margin": 2.5 / 2.54,
        "right_margin": 2.5 / 2.54,
        "bottom_margin": 1.5 / 2.54,
    }
    for attr, want in expected.items():
        got = getattr(sec, attr).inches
        if abs(got - want) > 0.02:
            errors.append(f"{attr}={got:.3f}, expected {want:.3f}")

    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    for prefix in EXPECTED[chapter]:
        if not any(h.startswith(prefix + " ") for h in headings):
            errors.append(f"missing section {prefix}")

    text = "\n".join(p.text for p in doc.paragraphs)

    caption_pattern = re.compile(rf"^Figure {chapter}\.\d+\s+[—-]\s+")
    captions = [p.text for p in doc.paragraphs if caption_pattern.match(p.text)]
    if len(captions) != 2:
        errors.append(f"figure captions={len(captions)}, expected 2")
    for i in (1, 2):
        if text.count(f"Figure {chapter}.{i}") < 2:
            errors.append(f"Figure {chapter}.{i} not referenced")

    words = count_words(doc)
    minimum = MIN_WORDS.get(chapter, 2300)
    if words < minimum:
        errors.append(f"word count too low: {words}")

    with ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        if len(media) != 2:
            errors.append(f"embedded images={len(media)}, expected 2")
        if xml.count("descr=\"Figure ") < 2:
            errors.append("missing figure alt descriptions")
        if "Times New Roman" not in zf.read("word/styles.xml").decode("utf-8"):
            errors.append("Times New Roman missing from styles")
        if "Courier New" not in xml:
            errors.append("Courier New code formatting missing")

    print(
        f"Chapter {chapter}: words={words:,} headings={len(headings)} "
        f"tables={len(doc.tables)} figures={len(captions)} size={path.stat().st_size:,} bytes"
    )
    return errors


def main() -> None:
    failures = []
    requested = sys.argv[1:]
    selected = {}
    if requested:
        for raw in requested:
            key = int(raw) if raw.isdigit() else raw.upper()
            if key not in FILES:
                failures.append(f"Unknown chapter: {raw}")
                continue
            selected[key] = FILES[key]
    else:
        selected = FILES

    for chapter, path in selected.items():
        if not path.exists():
            failures.append(f"Chapter {chapter}: missing file")
            continue
        failures.extend(f"Chapter {chapter}: {msg}" for msg in audit(chapter, path))
    if failures:
        print("\nFAILURES")
        print("\n".join(f"- {item}" for item in failures))
        raise SystemExit(1)
    print("\nALL CHAPTER AUDITS PASSED")


if __name__ == "__main__":
    main()
